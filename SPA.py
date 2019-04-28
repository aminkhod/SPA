#!/usr/bin/env python
# coding: utf-8

# In[1]:


from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.shared import Pt
import pandas as pd
from docx.shared import RGBColor
import numpy as np


# In[2]:


track = pd.read_csv('Tracking Sheet output.csv',encoding='latin1')
track


# In[3]:


# client = pd.read_csv('2018 ABACUS CLIENTS.csv')
# client.plot()


# In[4]:


# def lcs(X , Y): 
#     # find the length of the strings 
#     m = len(X) 
#     n = len(Y) 
  
#     # declaring the array for storing the dp values 
#     L = [[None]*(n+1) for i in range(m+1)] 
  
#     """Following steps build L[m+1][n+1] in bottom up fashion 
#     Note: L[i][j] contains length of LCS of X[0..i-1] 
#     and Y[0..j-1]"""
#     for i in range(m+1): 
#         for j in range(n+1): 
#             if i == 0 or j == 0 : 
#                 L[i][j] = 0
#             elif X[i-1] == Y[j-1]: 
#                 L[i][j] = L[i-1][j-1]+1
#             else: 
#                 L[i][j] = max(L[i-1][j] , L[i][j-1]) 
#     return L[m][n] 
# #end of function lcs 


# In[5]:


# def getClientRef(c, clientlist):
#     num = str(clientlist.loc[c,'SR. NO.']) + '.'
#     for char in clientlist.loc[c,'LIST OF COMPANY']:
#         if char != '.':
#             num = num + char
#         else:
#             break
#     return num
    
# def bestmatch(candid,clientlist):
#     # Remove all bad charachters.
#     candid = ''.join(e for e in candid if e.isalnum())
#     c = 0
#     maxAl = 0
#     candidRow = []
#     num = ''
#     maxLCStoClient = 0
#     for client in clientlist['LIST OF COMPANY']:

#         candidLcs = lcs(candid.lower(), client.lower())
#         LCStoClient = candidLcs / len(client)
#         if (candidLcs == len(candid)):
#             if (maxLCStoClient < LCStoClient):
#                 candidRow = list(clientlist.iloc[c])
#                 num = getClientRef(c, clientlist)
#                 maxLCStoClient = LCStoClient
#         c+=1


#     return num, candidRow
# bestmatch('ANA GENERAL TRADING LLC',client)


# In[6]:


# #Finding major company of REP.,
# def majorCoOfREP(candid,client):,
#     bestmatch(candid,client),
#     return ''


# In[7]:


def ordinal(date):
    if date == 1:
        ordin = 'ST'
    elif date == 2:
        ordin = 'ND'
    elif date == 3:
        ordin = 'RD'
    else:
        ordin = 'TH'
    return ordin


# In[8]:


def amontformat(totamount):
    totamount = str(totamount)
    c = len(totamount) - 1
    buf = ''
    buf1 = ''
    for i in range(c+1):
        if totamount[c - i] == '.':
            buf = totamount[c - i] + buf
            try:
                totamount1 = totamount.replace(buf,'')
#                 print(buf , totamount1)
            except:
                1+1
            if len(buf)<2:
                buf = buf[1:]
                buf = '.' +'00'
            elif len(buf) < 3:
                buf = buf[1:]
                buf = '.' + buf +'0'
            buf =  buf[:3]
            c1 = len(totamount1) -1
            count = 1
            
            for j in range(c1+1):
                if (count%3 == 0) and (j!=c1) :
                    buf1 = ',' + totamount1[c1 - j] + buf1
                else:
                    buf1 = totamount1[c1 - j] + buf1
                count += 1
#             print(buf1)
            break
        else:
            buf = totamount[c - i] + buf
    return buf1 + buf


# In[9]:


# amontformat(877000.0)


# In[10]:


def produceDocxFile(num,ref):
    #Title
    document =  Document('SPA.docx')
    paragraph = document.add_paragraph('\n \n \n')
    par_format = paragraph.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run= paragraph.add_run('PURCHASE AND SALE AGREEMENT')
    run.bold = True
    font = run.font
    font.name = 'Arial'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Nirmala UI'
    font.size = Pt(9)
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)

    paragraph = document.add_paragraph()
    par_format = paragraph.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run= paragraph.add_run('BY AND BETWEEN')
    run.bold = True

    #buyer
    paragraph = document.add_paragraph()
    par_format = paragraph.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run= paragraph.add_run(str(buyerName)+' - '+buyerCity)
    run.bold = True

    #repBuyer
    paragraph = document.add_paragraph()
    par_format = paragraph.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1= paragraph.add_run('Represented by: ')

    run= paragraph.add_run(str(buyerRep)+' - '+buyerRepCity)
    run.bold = True

    paragraph = document.add_paragraph()
    par_format = paragraph.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run= paragraph.add_run('AND')
    run.bold = True

    #source
    paragraph = document.add_paragraph()
    par_format = paragraph.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run= paragraph.add_run( sellerName+' - '+sellerCity)
    run.bold = True

    #rep source
    if (str(sellerRep) != str(np.nan)):
    #     print(np.nan, sellerRep)
        paragraph = document.add_paragraph()
        par_format = paragraph.paragraph_format
        par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1= paragraph.add_run('Represented by: ')
#         run1.bold = True
        run= paragraph.add_run(str(sellerRep)+' - '+str(sellerRepCity))
        run.bold = True


    #THIS PURCHASE AND SALE AGREEMENT is entered into this 8TH day of December 2018, by and between 
    paragraph = document.add_paragraph()
    par_format = paragraph.paragraph_format
    par_format.first_line_indent = Inches(0.27)
    run= paragraph.add_run('THIS PURCHASE AND SALE AGREEMENT is entered into this '+adate+', by and between ')

    run1= paragraph.add_run(str(buyerName)+' - '+buyerCity+ ' ')
    run1.bold = True

    run2= paragraph.add_run('Represented by: ')

    run3= paragraph.add_run(str(buyerRep)+' - '+buyerRepCity+ ' ')
    run3.bold = True

    run4= paragraph.add_run('(hereinafter referred as "Buyer") with office address at the United Arab Emirates and ')

    run5= paragraph.add_run(sellerName+' - '+sellerCity)
    run5.bold = True

    if (str(sellerRep) != str(np.nan)):
        run6= paragraph.add_run(' Represented by: ')

        run7= paragraph.add_run(str(sellerRep)+' - '+str(sellerRepCity))
        run7.bold = True
    run4= paragraph.add_run(' (hereinafter referred as "Seller”) with office address at the United Arab Emirates. ')

    #RECITALS:
    paragraph = document.add_paragraph()
    par_format = paragraph.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run= paragraph.add_run('RECITALS:')
    run.bold = True

    #WHEREAS, the SAUREX SPINNING SOULUTION GMBH & CO.KG - GILCHING, GERMANY and
    # SHANGHAI SEVEN BENEVOLENCE INTERNATIONAL TRADE CO. LTD - SHANGHAI, CHINA represented 
    #BEYZADE GENERAL TRADING L.L.C - DUBAI, U.A.E can enter into this Sale and Purchase Agreement
    #and sign pertinent documents with full rights under terms and conditions specified therein;
    paragraph = document.add_paragraph()
    par_format = paragraph.paragraph_format
    par_format.first_line_indent = Inches(0.27)
    run0= paragraph.add_run('WHEREAS')
    run0.bold = True

    run1= paragraph.add_run(', the ')

    run4= paragraph.add_run(sellerName+' - '+sellerCity )
    run4.bold = True

    if (str(sellerRep) != str(np.nan)):
        run2= paragraph.add_run(' Represented by: ')

        run3= paragraph.add_run(str(sellerRep)+' - '+str(sellerRepCity) + ' and ')
        run3.bold = True

    run6= paragraph.add_run(str(buyerName)+' - '+buyerCity+ ' ')
    run6.bold = True

    run7= paragraph.add_run('Represented by: ')

    run8= paragraph.add_run(str(buyerRep)+' - '+buyerRepCity+ ' ')
    run8.bold = True

    run7= paragraph.add_run(' can enter into this Sale and Purchase Agreement and sign pertinent documents with full ' +
                            'rights under terms and conditions specified therein;')
    font = run7.font
    font.name = 'Nirmala UI'
    font.size = Pt(9)
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)

    #WHEREAS, the SAUREX SPINNING SOULUTION GMBH & CO.KG - GILCHING, GERMANY desires to sell the Products
    #defined below and the SHANGHAI SEVEN BENEVOLENCE INTERNATIONAL TRADE CO. LTD - SHANGHAI, CHINA desires
    #to purchase the Products from SAUREX SPINNING SOULUTION GMBH & CO.KG - GILCHING, GERMANY.
    paragraph = document.add_paragraph()
    par_format = paragraph.paragraph_format
    par_format.first_line_indent = Inches(0.27)
    run0= paragraph.add_run('WHEREAS')
    run0.bold = True

    run1= paragraph.add_run(', the ')

    run2= paragraph.add_run(sellerName+' - '+sellerCity )
    run2.bold = True

    run3= paragraph.add_run(' desires to sell the Products defined below and the ')

    run4= paragraph.add_run(str(buyerName)+' - '+buyerCity+ ' ')
    run4.bold = True

    run5= paragraph.add_run(' desires to purchase the Products from ')

    run6= paragraph.add_run(sellerName+' - '+sellerCity + '. ')
    run6.bold = True
    font = run6.font
    font.name = 'Nirmala UI'
    font.size = Pt(9)
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)

    #NOW THEREFORE, in consideration of the mutual covenants and the agreements herein contained and other
    #goods and valuable (the receipt and sufficiency of which are hereby acknowledged) the parties agree as follows:
    paragraph = document.add_paragraph()
    par_format = paragraph.paragraph_format
    par_format.first_line_indent = Inches(0.27)
    run0= paragraph.add_run('NOW THEREFORE')
    run0.bold = True

    run1= paragraph.add_run(', in consideration of the mutual covenants and the agreements herein contained and other '
    + 'goods and valuable (the receipt and sufficiency of which are hereby acknowledged) the parties agree as follows: ')
    font = run1.font
    font.name = 'Nirmala UI'
    font.size = Pt(9)
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)

    #Sale of Product.  SAUREX SPINNING SOULUTION GMBH & CO.KG - GILCHING, GERMANY hereby sells to
    # SHANGHAI SEVEN BENEVOLENCE INTERNATIONAL TRADE CO. LTD - SHANGHAI, CHINA and 
    # SHANGHAI SEVEN BENEVOLENCE INTERNATIONAL TRADE CO. LTD - SHANGHAI, CHINA hereby purchases from
    # SAUREX SPINNING SOULUTION GMBH & CO.KG - GILCHING, GERMANY the product details below:
    paragraph = document.add_paragraph()
    styles = document.styles
    paragraph.style = document.styles['List Paragraph']

    par_format = paragraph.paragraph_format

    run0= paragraph.add_run('Sale of Product.')
    run0.bold = True
    run0.underline = True

    run1 = paragraph.add_run(' ' +sellerName+' - '+sellerCity)
    run1.bold = True

    run2= paragraph.add_run(' hereby sells to ' )

    run4= paragraph.add_run(str(buyerName)+' - '+buyerCity+ ' ')
    run4.bold = True

    run5= paragraph.add_run(' and ' )

    run6= paragraph.add_run(str(buyerName)+' - '+buyerCity+ ' ')
    run6.bold = True

    run7= paragraph.add_run(' hereby purchases from ' )

    run8= paragraph.add_run( sellerName+' - '+sellerCity )
    run8.bold = True

    run9= paragraph.add_run(' the product details below: ' )

    # p = document.add_paragraph('A plain paragraph having some ')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    records = []
    rec = []
    for i in range(len(pname)):
        rec.append(i+1)
        rec.append(pname[i])
        rec.append(pqty[i])
        rec.append(punit[i])
        rec.append(punitp[i])
        rec.append(ptotal[i])
        records.append(rec)
        rec = []

    # document.styles['OR']
    table = document.add_table(rows=1, cols=6, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'SR. NO.'
    hdr_cells[1].text = 'DESCRIPTION'
    hdr_cells[2].text = 'QTY'
    hdr_cells[3].text = 'UNIT OF MEASUREMENT'
    hdr_cells[4].text = 'UNIT PRICE'
    hdr_cells[5].text = 'AMOUNT'

    c = 0
    for SR, DESCRIPTION, UNITOFM, qty, UNITPRICE, AMOUNT in records:
        if c ==1:
            paragraph = document.add_paragraph('\n \n \n \n \n')
            table = document.add_table(rows=1, cols=6, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'SR. NO.'
            hdr_cells[1].text = 'DESCRIPTION'
            hdr_cells[2].text = 'QTY'
            hdr_cells[3].text = 'UNIT OF MEASUREMENT'
            hdr_cells[4].text = 'UNIT PRICE'
            hdr_cells[5].text = 'AMOUNT'

        c += 1 
        row_cells = table.add_row().cells
        row_cells[0].text = str(SR)
        row_cells[1].text = DESCRIPTION
        row_cells[2].text = str(UNITOFM)
        row_cells[3].text = qty
        row_cells[4].text = str(currency + ' ' + UNITPRICE)
        row_cells[5].text = str(currency + ' ' + str(AMOUNT))
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                par_format = paragraph.paragraph_format
                par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER            
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Arial Narrow'
                    font.size= Pt(10)
                    font.color.rgb = RGBColor(0x0, 0x0, 0x0)
    for cell in table.rows[0].cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                run.bold = True
    if len(records) == 1:
        paragraph = document.add_paragraph('\n \n \n')

    #2. Purchase Price.   SHANGHAI SEVEN BENEVOLENCE INTERNATIONAL TRADE CO. LTD - SHANGHAI, CHINA shall pay
    #to SAUREX SPINNING SOULUTION GMBH & CO.KG - GILCHING, GERMANY for the Products and for all obligations
    #specified herein, as full and complete consideration therefore, the sum of €      88,474.00 
    #(EIGHTY EIGHT THOUSAND FOUR HUNDRED SEVENTY FOUR EUROS ONLY).
    paragraph = document.add_paragraph()
    styles = document.styles
    paragraph.style = document.styles['List Paragraph']

    par_format = paragraph.paragraph_format

    run0= paragraph.add_run('Purchase Price.')
    run0.bold = True
    run0.underline = True

    run1 = paragraph.add_run(' ' + str(buyerName)+' - '+buyerCity+ ' ')
    run1.bold = True

    run2= paragraph.add_run('shall pay to ' )

    run4= paragraph.add_run( sellerName+' - '+sellerCity )
    run4.bold = True

    run5= paragraph.add_run(' for the Products and for all obligations specified herein, as full ' + 
                            'and complete consideration therefore, the sum of ' )

    run6= paragraph.add_run(currency + ' ' + str(totamount))
    run6.bold = True

    run7= paragraph.add_run(' (' + amountToWord + ').')


    #3.Payment.  Payment of the Purchase Price shall be made by 
    #SHANGHAI SEVEN BENEVOLENCE INTERNATIONAL TRADE CO. LTD - SHANGHAI, CHINA or its representative 
    #BEYZADE GENERAL TRADING L.L.C - DUBAI, U.A.E to SAUREX SPINNING SOULUTION GMBH & CO.KG - GILCHING, GERMANY 
    #in full payment in advance before the delivery date.
    paragraph = document.add_paragraph()
    styles = document.styles
    paragraph.style = document.styles['List Paragraph']

    par_format = paragraph.paragraph_format

    run0= paragraph.add_run('Payment.')
    run0.bold = True
    run0.underline = True

    run1= paragraph.add_run(' Payment of the Purchase Price shall be made by ' )

    run2= paragraph.add_run( str(buyerName)+' - '+buyerCity )
    run2.bold = True

    run3= paragraph.add_run(' or its representative ')
    run3.bold = False

    run4= paragraph.add_run( str(buyerRep)+' - '+buyerRepCity )
    run4.bold = True

    run5= paragraph.add_run(' to ')

    run6= paragraph.add_run( sellerName+' - '+sellerCity )
    run6.bold = True

    if (str(sellerRep) != str(np.nan)):
        run7= paragraph.add_run(' or its representative  ' )
        run8= paragraph.add_run(str(sellerRep)+' - '+str(sellerRepCity))
        run8.bold = True
    run9= paragraph.add_run(' in full payment in advance before the delivery date.')


    #Acceptance.   “Acceptance" of the Product shall be deemed to occur on the date when, in the reasonable
    #opinion of SHANGHAI SEVEN BENEVOLENCE INTERNATIONAL TRADE CO. LTD - SHANGHAI, CHINA the Product conforms 
    #to the Specifications, and has continuously operated in compliance with the Specifications for thirty (30) 
    #days after Product Turnover
    paragraph = document.add_paragraph()
    styles = document.styles
    paragraph.style = document.styles['List Paragraph']

    par_format = paragraph.paragraph_format

    run0= paragraph.add_run('Acceptance.')
    run0.bold = True
    run0.underline = True

    run1= paragraph.add_run(' “Acceptance" of the Product shall be deemed to occur on the date when, '+
                            'in the reasonable opinion of ' )

    run2= paragraph.add_run(str(buyerName)+' - '+buyerCity)
    run2.bold = True

    run4= paragraph.add_run(' the Product conforms to the Specifications, and has continuously operated in compliance ' + 
                            'with the Specifications for thirty (30) days after Product Turnover.' )

    #Indemnification. In the event either party breaches or is deemed to have breached any of the 
    #representations and warranties contained in this Agreement, or fails to perform or comply with any
    #of the covenants and agreements set forth in this Agreement, it shall hold harmless, indemnify and
    #defend the other party, and its directors, officers, shareholders, attorneys, representatives and
    #agents, from and against any damages incurred by the non-defaulting party. 
    paragraph = document.add_paragraph()
    styles = document.styles
    paragraph.style = document.styles['List Paragraph']

    par_format = paragraph.paragraph_format

    run0= paragraph.add_run('Indemnification.')
    run0.bold = True
    run0.underline = True

    run1= paragraph.add_run(' In the event either party breaches or is deemed to have breached any of the '+
    'representations and warranties contained in this Agreement, or fails to perform or comply with any'+
    ' of the covenants and agreements set forth in this Agreement, it shall hold harmless, indemnify and'+
    ' defend the other party, and its directors, officers, shareholders, attorneys, representatives and'+
    ' agents, from and against any damages incurred by the non-defaulting party.')


    #General. SAUREX SPINNING SOULUTION GMBH & CO.KG - GILCHING, GERMANY shall perform this Agreement
    #in compliance with all applicable local laws, rules, regulations, and ordinances, and represents that
    #it shall have obtained all licenses and permits required by law to engage in the activities necessary
    #to perform its obligations under this Agreement.  
    paragraph = document.add_paragraph()
    styles = document.styles
    paragraph.style = document.styles['List Paragraph']

    par_format = paragraph.paragraph_format

    run0= paragraph.add_run('General.')
    run0.bold = True
    run0.underline = True

    run2 = paragraph.add_run(' ' + sellerName+' - '+sellerCity + ' ')
    run2.bold = True

    run1= paragraph.add_run('shall perform this Agreement in compliance with all applicable local laws'+
                            ', rules, regulations, and ordinances, and represents that it shall have ' + 
                            'obtained all licenses and permits required by law to engage in the activities ' +
                            'necessary to perform its obligations under this Agreement. ' )
    s = ''
    for i in range(7-len(records)):
        s += '\n'
    paragraph = document.add_paragraph(s)
    paragraph = document.add_paragraph(int(1.33 * len(sellerName+' - '+sellerCity))*'_' + '\n')
    run0= paragraph.add_run(sellerName+' - '+sellerCity)
    run0.bold = True

    # document.add_picture('sign.png', width=Inches(2))
    paragraph = document.add_paragraph('\n \n \n \n')
    paragraph = document.add_paragraph(int(1.33 * len(str(buyerName)+' - '+buyerCity)) * '_'+'\n')
    run0= paragraph.add_run(str(buyerName)+' - '+buyerCity)
    run0.bold = True

    document.save('output/' + str(ref) + '.docx')


# In[11]:


#PO tracking sheet
vect = []
for co in range(len(track['REF'])):
    vector = list(track.loc[co])
    sellerRep = vector[1]
    sellerRepCity = vector[2]
    try:
        sellerRepCity = sellerRepCity.replace('-',',')
    except:
        1+1
    try:
        sellerRepCity = sellerRepCity.lower()
        sellerRepCity = sellerRepCity.replace('uae','U.A.E')
        sellerRepCity = sellerRepCity.upper()
    except:
        1+1

    sellerName = vector[3]
    sellerCity = vector[4]
    try:
        sellerCity = sellerCity.replace('-',',')
    except:
        1+1
    try:
        sellerCity = sellerCity.lower()
        sellerCity = sellerCity.replace('uae','U.A.E')
        sellerCity = sellerCity.upper()
    except:
        1+1

    buyerName = vector[5]
    buyerCity = vector[6]
    try:
        buyerCity = buyerCity.replace('-',',')
    except:
        1+1
    try:
        buyerCity = buyerCity.lower()
        buyerCity = buyerCity.replace('uae','U.A.E')
        buyerCity = buyerCity.upper()
    except:
        1+1
    buyerRep = vector[7]
    buyerRepCity = vector[8]
    try:
        buyerRepCity = buyerRepCity.replace('-',',')
    except:
        1+1
    try:
        buyerRepCity = buyerRepCity.lower()
        buyerRepCity = buyerRepCity.replace('uae','U.A.E')
        buyerRepCity = buyerRepCity.upper()
    except:
        1+1
    fdate = vector[41]
#     print(fdate)
    adate = ''
    for ch in fdate:
        if ch !='-':
            adate = adate + ch
        else:
            try:
                adate = adate + ordinal(int(adate)) + ' day of '
            except:
                adate = adate + ' '
    product = vector[14:39]
    product
    pname = []
    pqty = []
    punit = []
    punitp = []
    ptotal = []
    cur = []
    for i in range(5):
        if not str(product[i*5]) == str(np.nan):
            pname.append(product[i*5])
            pqty.append(product[i*5 +1])
            punit.append(product[i*5+2])
            punitp.append(amontformat(str(product[i*5+3])))
            ptotal.append(amontformat(str(product[i*5+4])))
    currency = vector[40]
    amount = vector[12]
    #money format
    totamount = amontformat(str(amount))
    amountToWord = vector[13]
    if(float(amount) >= 54449.00):
        produceDocxFile(co, track.-loc[co, 'REF'])
        


# In[ ]:




